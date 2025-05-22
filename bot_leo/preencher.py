from docx import Document

def substituir_em_runs(paragrafos, dados):
    for paragrafo in paragrafos:
        for run in paragrafo.runs:
            for chave, valor in dados.items():
                if f'{{{{{chave}}}}}' in run.text:
                    run.text = run.text.replace(f'{{{{{chave}}}}}', valor)

def preencher_modelo(caminho_modelo, caminho_saida, dados):
    doc = Document(caminho_modelo)

    # Substitui nos parágrafos principais
    substituir_em_runs(doc.paragraphs, dados)

    # Substitui nas tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                substituir_em_runs(celula.paragraphs, dados)

    # Substitui no cabeçalho e rodapé
    for section in doc.sections:
        for part in [section.header, section.footer]:
            substituir_em_runs(part.paragraphs, dados)
           
    doc.save(caminho_saida)
